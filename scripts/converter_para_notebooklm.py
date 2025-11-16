# -*- coding: utf-8 -*-
"""
Script para converter documentos .md para formatos aceitos pelo NotebookLM
NotebookLM aceita: PDF, DOCX, TXT, e alguns outros formatos
"""

import os
import sys
from pathlib import Path
import subprocess

project_root = Path(__file__).parent.parent
obsidian_path = project_root / "Obsidian_guardar aqui"
notebooklm_path = project_root / "NotebookLM"

def verificar_dependencias():
    """Verifica se as dependências necessárias estão instaladas."""
    dependencias = {
        "pypandoc": "pypandoc",
        "markdown": "markdown",
        "pdfkit": "pdfkit",
        "docx": "python-docx",
    }
    
    faltando = []
    for nome, pacote in dependencias.items():
        try:
            __import__(nome)
        except ImportError:
            faltando.append(pacote)
    
    return faltando

def instalar_dependencias():
    """Instala dependências necessárias."""
    print("Verificando dependencias...")
    faltando = verificar_dependencias()
    
    if faltando:
        print(f"Dependencias faltando: {', '.join(faltando)}")
        print("Instalando...")
        for pacote in faltando:
            subprocess.run([sys.executable, "-m", "pip", "install", pacote], check=False)
        print("Dependencias instaladas!")
    else:
        print("Todas as dependencias estao instaladas!")

def converter_md_para_txt(origem: Path, destino: Path):
    """Converte .md para .txt (formato simples, sempre funciona)."""
    try:
        with open(origem, 'r', encoding='utf-8') as f:
            conteudo = f.read()
        
        # Remove formatação markdown básica (opcional)
        # Pode manter o markdown também, TXT aceita
        destino.parent.mkdir(parents=True, exist_ok=True)
        
        with open(destino, 'w', encoding='utf-8') as f:
            f.write(conteudo)
        
        return True
    except Exception as e:
        print(f"Erro ao converter {origem.name} para TXT: {e}")
        return False

def converter_md_para_docx(origem: Path, destino: Path):
    """Converte .md para .docx usando python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt
        import re
        
        doc = Document()
        
        with open(origem, 'r', encoding='utf-8') as f:
            conteudo = f.read()
        
        linhas = conteudo.split('\n')
        em_codigo = False
        codigo_linhas = []
        
        for linha in linhas:
            linha_original = linha
            linha = linha.rstrip()
            
            # Blocos de código
            if linha.startswith('```'):
                if em_codigo:
                    # Fecha bloco de código
                    if codigo_linhas:
                        para = doc.add_paragraph('\n'.join(codigo_linhas))
                        para.style = 'No Spacing'
                        run = para.runs[0] if para.runs else para.add_run()
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                    codigo_linhas = []
                    em_codigo = False
                else:
                    em_codigo = True
                continue
            
            if em_codigo:
                codigo_linhas.append(linha)
                continue
            
            # Títulos
            if linha.startswith('# '):
                doc.add_heading(linha[2:], level=1)
            elif linha.startswith('## '):
                doc.add_heading(linha[3:], level=2)
            elif linha.startswith('### '):
                doc.add_heading(linha[4:], level=3)
            elif linha.startswith('#### '):
                doc.add_heading(linha[5:], level=4)
            elif linha.startswith('##### '):
                doc.add_heading(linha[6:], level=5)
            elif linha.startswith('###### '):
                doc.add_heading(linha[7:], level=6)
            # Listas
            elif linha.startswith('- ') or linha.startswith('* '):
                doc.add_paragraph(linha[2:], style='List Bullet')
            elif linha.startswith('1. ') or re.match(r'^\d+\. ', linha):
                doc.add_paragraph(re.sub(r'^\d+\. ', '', linha), style='List Number')
            # Links
            elif re.search(r'\[([^\]]+)\]\(([^\)]+)\)', linha):
                para = doc.add_paragraph()
                for match in re.finditer(r'\[([^\]]+)\]\(([^\)]+)\)', linha):
                    texto_link = match.group(1)
                    url = match.group(2)
                    run = para.add_run(texto_link)
                    run.hyperlink.address = url
                    run.font.color.rgb = None  # Cor padrão de link
            # Parágrafo normal
            elif linha.strip():
                doc.add_paragraph(linha)
            else:
                doc.add_paragraph()  # Linha vazia
        
        destino.parent.mkdir(parents=True, exist_ok=True)
        doc.save(destino)
        return True
    except ImportError:
        print("python-docx nao instalado. Instale com: pip install python-docx")
        return False
    except Exception as e:
        print(f"Erro ao converter {origem.name} para DOCX: {e}")
        return False

def converter_md_para_pdf(origem: Path, destino: Path):
    """Converte .md para .pdf usando markdown + weasyprint ou pdfkit."""
    try:
        import markdown
        from weasyprint import HTML
        
        with open(origem, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Converte markdown para HTML
        html_content = markdown.markdown(md_content, extensions=['extra', 'codehilite'])
        
        # Adiciona estilo básico
        html_full = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
                h1 {{ color: #333; border-bottom: 2px solid #333; }}
                h2 {{ color: #555; margin-top: 30px; }}
                code {{ background: #f4f4f4; padding: 2px 5px; }}
                pre {{ background: #f4f4f4; padding: 10px; overflow-x: auto; }}
            </style>
        </head>
        <body>
        {html_content}
        </body>
        </html>
        """
        
        destino.parent.mkdir(parents=True, exist_ok=True)
        HTML(string=html_full).write_pdf(destino)
        return True
    except ImportError:
        return False
    except Exception as e:
        print(f"Erro ao converter {origem.name} para PDF: {e}")
        return False

def converter_todos_arquivos(formato='txt'):
    """Converte todos os arquivos .md para o formato especificado."""
    
    if not notebooklm_path.exists():
        print("Pasta NotebookLM nao existe. Execute preparar_para_notebooklm.py primeiro.")
        return
    
    print("=" * 70)
    print(f"CONVERTENDO ARQUIVOS PARA {formato.upper()}")
    print("=" * 70)
    print()
    
    # Mapeamento de pastas
    mapeamento = {
        "01-Fundamentos": [
            "PROJETO-IA-TEST.md",
            "00-MAPA-DE-AGENTES.md",
            "ESTRUTURA-PROJETO.md",
            "00-ERROS-E-CONFIGURACOES-PENDENTES.md",
        ],
        "02-LangChain-LangGraph": [
            "LANGCHAIN-LANGGRAPH-GUIA.md",
            "LANGCHAIN-FUNDAMENTOS.md",
            "LANGGRAPH-CONCEITOS.md",
            "LANGGRAPH-WORKFLOWS.md",
            "LANGCHAIN-NEO4J.md",
            "LANGGRAPH-PADROES.md",
            "LANGGRAPH-AGENTES.md",
            "LANGCHAIN-EXEMPLOS.md",
            "PREPARACAO-LANGCHAIN.md",
            "RESUMO-LANGCHAIN-PREPARACAO.md",
        ],
        "03-Agentes": [
            "Agentes/Orchestrator.md",
            "Agentes/System-Health.md",
            "Agentes/DB-Manager.md",
            "Agentes/MCP-Manager.md",
            "Agentes/Neo4j-GraphRAG.md",
            "Agentes/Obsidian-Integration.md",
            "Agentes/Kestra-Agent.md",
            "Agentes/Docker-Integration.md",
            "Agentes/Git-Integration.md",
        ],
        "04-Configuracao": [
            "CONFIGURACOES-APLICADAS.md",
            "NEO4J-CONFIGURADO.md",
            "COMO-CONFIGURAR-NEO4J-URI.md",
            "ANALISE-BANCOS-DADOS.md",
        ],
        "05-Exemplos": [
            "VIDEOS_MCP_AGENTES.md",
            "OTIMIZACAO_AGENTES.md",
        ],
        "06-Referencias": [
            "README_ESTRUTURA.md",
            "RESUMO-MAPA-AGENTES.md",
        ]
    }
    
    convertidos = 0
    erros = 0
    
    for pasta_destino, arquivos in mapeamento.items():
        destino_pasta = notebooklm_path / pasta_destino
        
        for arquivo_nome in arquivos:
            origem = obsidian_path / arquivo_nome
            
            if not origem.exists():
                continue
            
            # Define extensão do destino
            if formato == 'txt':
                destino = destino_pasta / origem.name.replace('.md', '.txt')
                sucesso = converter_md_para_txt(origem, destino)
            elif formato == 'docx':
                destino = destino_pasta / origem.name.replace('.md', '.docx')
                sucesso = converter_md_para_docx(origem, destino)
            elif formato == 'pdf':
                destino = destino_pasta / origem.name.replace('.md', '.pdf')
                sucesso = converter_md_para_pdf(origem, destino)
            else:
                print(f"Formato nao suportado: {formato}")
                return
            
            if sucesso:
                convertidos += 1
                print(f"Convertido: {origem.name} -> {destino.name}")
            else:
                erros += 1
                print(f"Erro: {origem.name}")
    
    print()
    print("=" * 70)
    print(f"CONVERSAO CONCLUIDA!")
    print(f"  Convertidos: {convertidos}")
    print(f"  Erros: {erros}")
    print("=" * 70)
    print()
    print(f"Arquivos convertidos para {formato.upper()} na pasta NotebookLM/")
    print("Agora voce pode adicionar esses arquivos no NotebookLM!")

def main():
    """Função principal."""
    import argparse
    
    parser = argparse.ArgumentParser(description='Converter arquivos .md para formatos do NotebookLM')
    parser.add_argument('--formato', choices=['txt', 'docx', 'pdf'], default='txt',
                       help='Formato de saída (txt, docx, pdf)')
    parser.add_argument('--instalar', action='store_true',
                       help='Instalar dependências necessárias')
    
    args = parser.parse_args()
    
    if args.instalar:
        instalar_dependencias()
        return
    
    # Verifica dependências básicas
    if args.formato == 'txt':
        # TXT não precisa de dependências extras
        pass
    elif args.formato == 'docx':
        faltando = verificar_dependencias()
        if 'python-docx' in faltando:
            print("Instalando python-docx...")
            subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"], check=False)
    elif args.formato == 'pdf':
        faltando = verificar_dependencias()
        if 'markdown' in faltando or 'weasyprint' in faltando:
            print("Instalando dependências para PDF...")
            subprocess.run([sys.executable, "-m", "pip", "install", "markdown", "weasyprint"], check=False)
    
    converter_todos_arquivos(args.formato)

if __name__ == "__main__":
    main()

